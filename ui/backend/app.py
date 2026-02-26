from __future__ import annotations

import json
import io
import logging
import os
import re
import threading
import uuid
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, Optional
from urllib.parse import quote

from flask import (
    Flask,
    Response,
    jsonify,
    redirect,
    request,
    send_from_directory,
    stream_with_context,
)
from werkzeug.exceptions import HTTPException

from . import pipeline_manager as pm

LOGGER = logging.getLogger(__name__)

BASE_DIR = Path(__file__).resolve().parent
FRONTEND_DIR = BASE_DIR.parent / "frontend"
EXAMPLES_DIR = BASE_DIR.parent.parent / "examples"
KB_TASKS = {}
LLMS_DOC_PATH = BASE_DIR.parent.parent / "docs" / "llms.txt"
LLMS_DOC_CACHE = None
DOCX_MIME_TYPE = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)
MEMORY_ROOT = BASE_DIR.parent.parent / "data" / "user_memory"
DEFAULT_MEMORY_USER = "default"
MEMORY_FILENAME = "MEMORY.md"
MEMORY_DEFAULT_CONTENT = "# MEMORY\ni am jack. i like LLMs.\n"
MEMORY_USER_ID_PATTERN = re.compile(r"^[A-Za-z0-9_-]+$")


def load_llms_doc() -> str:
    """Load docs/llms.txt once and cache it for system prompt usage."""
    global LLMS_DOC_CACHE
    if LLMS_DOC_CACHE is not None:
        return LLMS_DOC_CACHE

    try:
        LLMS_DOC_CACHE = LLMS_DOC_PATH.read_text(encoding="utf-8")
        LOGGER.info("Loaded llms.txt reference (%d chars)", len(LLMS_DOC_CACHE))
    except FileNotFoundError:
        LOGGER.warning("llms.txt not found at %s", LLMS_DOC_PATH)
        LLMS_DOC_CACHE = ""
    except Exception as e:
        LOGGER.error("Failed to load llms.txt: %s", e)
        LLMS_DOC_CACHE = ""

    return LLMS_DOC_CACHE


def _normalize_memory_user_id(raw_user_id: Optional[str]) -> str:
    user_id = str(raw_user_id or DEFAULT_MEMORY_USER).strip() or DEFAULT_MEMORY_USER
    if not MEMORY_USER_ID_PATTERN.fullmatch(user_id):
        raise ValueError("Invalid user_id format")
    return user_id


def _ensure_memory_file(user_id: str) -> Path:
    user_dir = MEMORY_ROOT / user_id
    project_dir = user_dir / "project"
    project_dir.mkdir(parents=True, exist_ok=True)
    memory_path = user_dir / MEMORY_FILENAME
    if not memory_path.exists():
        memory_path.write_text(MEMORY_DEFAULT_CONTENT, encoding="utf-8")
    return memory_path


def _normalize_export_title(question_text: str) -> str:
    normalized = re.sub(r"\s+", " ", str(question_text or "")).strip()
    normalized = re.sub(r"^#+\s*", "", normalized)
    return normalized or "Chat Export"


def _sanitize_export_filename(title: str) -> str:
    safe = re.sub(r'[\\/:*?"<>|]', "", title)
    safe = re.sub(r"\s+", "-", safe)
    safe = re.sub(r"-+", "-", safe).strip(".-")
    return (safe[:80] or "chat-export").strip() or "chat-export"


def _ascii_fallback_filename(filename: str, default_basename: str = "chat-export") -> str:
    safe = str(filename or "").replace("\r", "").replace("\n", "")
    basename, ext = os.path.splitext(safe)
    ascii_basename = basename.encode("ascii", "ignore").decode("ascii")
    ascii_basename = re.sub(r"[^A-Za-z0-9._-]+", "-", ascii_basename)
    ascii_basename = re.sub(r"-+", "-", ascii_basename).strip("-.")
    if not ascii_basename:
        ascii_basename = default_basename
    ascii_ext = ext if re.fullmatch(r"\.[A-Za-z0-9]+", ext or "") else ".docx"
    return f"{ascii_basename}{ascii_ext}"


def _build_content_disposition(filename: str) -> str:
    safe = str(filename or "").replace("\r", "").replace("\n", "")
    ascii_filename = _ascii_fallback_filename(safe)
    utf8_filename = quote(safe, safe="")
    return (
        f"attachment; filename=\"{ascii_filename}\"; "
        f"filename*=UTF-8''{utf8_filename}"
    )


def _build_source_map(sources: Any) -> Dict[int, Dict[str, str]]:
    source_map: Dict[int, Dict[str, str]] = {}
    if not isinstance(sources, list):
        return source_map

    for src in sources:
        if not isinstance(src, dict):
            continue
        raw_ref_id = src.get("displayId") or src.get("id")
        try:
            ref_id = int(raw_ref_id)
        except (TypeError, ValueError):
            continue
        if ref_id in source_map:
            continue
        source_map[ref_id] = {
            "title": str(src.get("title") or "").strip(),
            "content": str(src.get("content") or ""),
        }
    return source_map


def _ordered_reference_ids(answer_text: str, source_map: Dict[int, Dict[str, str]]) -> list[int]:
    used_ids: list[int] = []
    for ref_text in re.findall(r"\[(\d+)\]", str(answer_text or "")):
        ref_id = int(ref_text)
        if ref_id not in used_ids:
            used_ids.append(ref_id)
    return used_ids if used_ids else sorted(source_map.keys())


def _set_run_fonts(
    run: Any,
    *,
    size_pt: Optional[float] = None,
    bold: Optional[bool] = None,
    mono: bool = False,
) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt

    latin_font = "Consolas" if mono else "Calibri"
    east_asia_font = "等线" if mono else "Microsoft YaHei"
    run.font.name = latin_font

    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:ascii"), latin_font)
    r_fonts.set(qn("w:hAnsi"), latin_font)
    r_fonts.set(qn("w:eastAsia"), east_asia_font)

    if size_pt is not None:
        run.font.size = Pt(size_pt)
    if bold is not None:
        run.bold = bold


def _strip_markdown_links(text: str) -> str:
    cleaned = re.sub(r"<a[^>]*>(.*?)</a>", r"\1", text, flags=re.IGNORECASE)
    cleaned = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r"\1 (\2)", cleaned)
    return cleaned


def _append_markdown_to_docx(document: Any, markdown_text: str) -> None:
    if not str(markdown_text or "").strip():
        return

    heading_size_rules = {
        1: 20,
        2: 16,
        3: 14,
        4: 13,
        5: 12,
        6: 11,
    }
    in_code_block = False

    for raw_line in str(markdown_text).splitlines():
        line = raw_line.rstrip("\r")
        stripped = line.strip()

        if stripped.startswith("```"):
            in_code_block = not in_code_block
            continue

        if re.match(r'^<a\s+id="[^"]+"\s*></a>$', stripped):
            continue

        if in_code_block:
            para = document.add_paragraph()
            run = para.add_run(line)
            _set_run_fonts(run, size_pt=10.5, mono=True)
            continue

        if not stripped:
            document.add_paragraph("")
            continue

        heading_match = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if heading_match:
            level = len(heading_match.group(1))
            heading_text = _strip_markdown_links(heading_match.group(2).strip())
            para = document.add_paragraph()
            run = para.add_run(heading_text)
            _set_run_fonts(run, size_pt=heading_size_rules.get(level, 11), bold=True)
            continue

        bullet_match = re.match(r"^[-*]\s+(.*)$", stripped)
        if bullet_match:
            para = document.add_paragraph(style="List Bullet")
            run = para.add_run(_strip_markdown_links(bullet_match.group(1).strip()))
            _set_run_fonts(run, size_pt=11)
            continue

        ordered_match = re.match(r"^\d+\.\s+(.*)$", stripped)
        if ordered_match:
            para = document.add_paragraph(style="List Number")
            run = para.add_run(_strip_markdown_links(ordered_match.group(1).strip()))
            _set_run_fonts(run, size_pt=11)
            continue

        para = document.add_paragraph()
        run = para.add_run(_strip_markdown_links(line))
        _set_run_fonts(run, size_pt=11)


def _build_chat_export_docx(
    question_text: str,
    answer_text: str,
    sources: Any,
) -> tuple[bytes, str]:
    try:
        from docx import Document
    except Exception as e:
        raise RuntimeError("python-docx is required for DOCX export.") from e

    export_title = _normalize_export_title(question_text)
    answer = str(answer_text or "")
    source_map = _build_source_map(sources)
    ordered_ref_ids = _ordered_reference_ids(answer, source_map)

    document = Document()

    title_para = document.add_paragraph()
    title_run = title_para.add_run(export_title)
    _set_run_fonts(title_run, size_pt=20, bold=True)

    document.add_paragraph("")

    answer_heading = document.add_paragraph()
    answer_heading_run = answer_heading.add_run("Answer")
    _set_run_fonts(answer_heading_run, size_pt=16, bold=True)

    _append_markdown_to_docx(document, answer if answer.strip() else "(empty)")

    if ordered_ref_ids:
        document.add_paragraph("")
        ref_heading = document.add_paragraph()
        ref_heading_run = ref_heading.add_run("References")
        _set_run_fonts(ref_heading_run, size_pt=16, bold=True)

        for ref_id in ordered_ref_ids:
            source = source_map.get(ref_id, {})
            title = source.get("title") or f"Reference {ref_id}"
            content = str(source.get("content") or "").strip()

            item_heading = document.add_paragraph()
            item_heading_run = item_heading.add_run(f"[{ref_id}] {title}")
            _set_run_fonts(item_heading_run, size_pt=14, bold=True)

            if content:
                _append_markdown_to_docx(document, content)
            else:
                para = document.add_paragraph()
                run = para.add_run("Reference content unavailable.")
                _set_run_fonts(run, size_pt=11)

    buffer = io.BytesIO()
    document.save(buffer)
    filename = (
        f"{_sanitize_export_filename(export_title)}-"
        f"{datetime.now().strftime('%Y%m%d-%H%M%S')}.docx"
    )
    return buffer.getvalue(), filename


def _run_kb_background(
    task_id: str,
    pipeline_name: str,
    target_file: str,
    output_dir: str,
    collection_name: str,
    index_mode: str,
    chunk_params: Optional[Dict[str, Any]] = None,
    embedding_params: Optional[Dict[str, Any]] = None,
) -> None:
    """Run knowledge base pipeline in background thread.

    Args:
        task_id: Unique task identifier
        pipeline_name: Name of the pipeline to run
        target_file: Path to target file
        output_dir: Output directory path
        collection_name: Milvus collection name
        index_mode: Index mode ("append" or "overwrite")
        chunk_params: Optional chunking parameters
        embedding_params: Optional embedding parameters
    """
    LOGGER.info(f"Task {task_id} started: {pipeline_name}")
    try:
        result = pm.run_kb_pipeline_tool(
            pipeline_name=pipeline_name,
            target_file_path=target_file,
            output_dir=output_dir,
            collection_name=collection_name,
            index_mode=index_mode,
            chunk_params=chunk_params,
            embedding_params=embedding_params,
        )
        KB_TASKS[task_id]["status"] = "success"
        KB_TASKS[task_id]["result"] = result
        KB_TASKS[task_id]["completed_at"] = datetime.now().isoformat()
        LOGGER.info(f"Task {task_id} completed successfully.")

    except Exception as e:
        LOGGER.error(f"Task {task_id} failed: {e}", exc_info=True)
        KB_TASKS[task_id]["status"] = "failed"
        KB_TASKS[task_id]["error"] = str(e)


def create_app(admin_mode: bool = False) -> Flask:
    """Create and configure Flask application.

    Args:
        admin_mode: Whether to run in admin mode (default: False)

    Returns:
        Configured Flask application instance
    """
    app = Flask(__name__, static_folder=str(FRONTEND_DIR), static_url_path="")
    app.config["ADMIN_MODE"] = admin_mode

    @app.errorhandler(pm.PipelineManagerError)
    def handle_pipeline_error(
        err: pm.PipelineManagerError,
    ) -> tuple[Dict[str, str], int]:
        """Handle pipeline manager errors.

        Args:
            err: Pipeline manager error

        Returns:
            JSON error response with 400 status
        """
        LOGGER.error(f"Pipeline error: {err}")
        return jsonify({"error": str(err)}), 400

    @app.errorhandler(Exception)
    def handle_generic_error(err: Exception) -> tuple[Dict[str, str], int]:
        """Handle generic exceptions.

        Args:
            err: Exception instance

        Returns:
            JSON error response with 500 status
        """
        if isinstance(err, HTTPException):
            return (
                jsonify(
                    {
                        "error": err.name,
                        "details": err.description,
                    }
                ),
                err.code or 500,
            )
        LOGGER.error(f"System error: {err}", exc_info=True)
        return jsonify({"error": "Internal Server Error", "details": str(err)}), 500

    @app.route("/favicon.svg")
    def favicon():
        return send_from_directory(
            os.path.join(app.static_folder), "favicon.svg", mimetype="image/svg+xml"
        )

    @app.route("/")
    def index() -> Response:
        """Serve frontend index page."""
        return send_from_directory(app.static_folder, "index.html")

    # Direct access to /chat or /settings also returns frontend entry,
    # facilitating separation of configuration/chat pages
    @app.route("/chat")
    def chat_page() -> Response:
        """Serve frontend chat page."""
        return send_from_directory(app.static_folder, "index.html")

    @app.route("/settings")
    def settings_page() -> Response:
        """Serve frontend settings page."""
        return send_from_directory(app.static_folder, "index.html")

    @app.route("/config")
    def config_page() -> Response:
        """Redirect legacy /config route to /settings."""
        return redirect("/settings")

    @app.route("/api/memory", methods=["GET", "PUT"])
    @app.route("/api/memory/<string:user_id>", methods=["GET", "PUT"])
    def memory_file(user_id: str = DEFAULT_MEMORY_USER) -> Response:
        """Read or update per-user MEMORY.md content."""
        try:
            normalized_user = _normalize_memory_user_id(user_id)
        except ValueError as e:
            return jsonify({"error": str(e)}), 400

        memory_path = _ensure_memory_file(normalized_user)
        relative_path = str(memory_path.relative_to(BASE_DIR.parent.parent))

        if request.method == "GET":
            try:
                content = memory_path.read_text(encoding="utf-8")
            except Exception as e:
                LOGGER.error(
                    "Failed to read memory file for user %s: %s",
                    normalized_user,
                    e,
                    exc_info=True,
                )
                return jsonify({"error": str(e)}), 500
            return jsonify(
                {
                    "user_id": normalized_user,
                    "path": relative_path,
                    "content": content,
                }
            )

        payload = request.get_json(force=True) or {}
        content = payload.get("content", "")
        if not isinstance(content, str):
            return jsonify({"error": "content must be a string"}), 400

        try:
            memory_path.write_text(content, encoding="utf-8")
        except Exception as e:
            LOGGER.error(
                "Failed to save memory file for user %s: %s",
                normalized_user,
                e,
                exc_info=True,
            )
            return jsonify({"error": str(e)}), 500

        return jsonify(
            {
                "status": "saved",
                "user_id": normalized_user,
                "path": relative_path,
            }
        )

    @app.route("/api/config/mode", methods=["GET"])
    def get_app_mode() -> Response:
        """Return the application mode (admin or chat-only).

        Returns:
            JSON response with admin_mode flag
        """
        return jsonify({"admin_mode": app.config.get("ADMIN_MODE", False)})

    @app.route("/api/templates", methods=["GET"])
    def list_templates() -> Response:
        """List available pipeline templates.

        Returns:
            JSON response with list of template names and contents
        """
        templates = []
        for f in sorted(EXAMPLES_DIR.glob("*.yaml")):
            try:
                templates.append(
                    {"name": f.stem, "content": f.read_text(encoding="utf-8")}
                )
            except Exception:
                continue
        return jsonify(templates)

    @app.route("/api/servers", methods=["GET"])
    def servers() -> Response:
        """List available MCP servers.

        Returns:
            JSON response with list of servers
        """
        return jsonify(pm.list_servers())

    @app.route("/api/tools", methods=["GET"])
    def tools() -> Response:
        """List available tools from all servers.

        Returns:
            JSON response with list of tools
        """
        return jsonify(
            [
                {
                    "id": tool.identifier,
                    "server": tool.server,
                    "tool": tool.tool,
                    "kind": tool.kind,
                    "input": tool.input_spec,
                    "output": tool.output_spec,
                }
                for tool in pm.list_server_tools()
            ]
        )

    @app.route("/api/pipelines", methods=["GET"])
    def list_pipelines() -> Response:
        """List available pipelines.

        Returns:
            JSON response with list of pipelines
        """
        return jsonify(pm.list_pipelines())

    @app.route("/api/pipelines", methods=["POST"])
    def save_pipeline() -> Response:
        """Save a new pipeline.

        Returns:
            JSON response with save result
        """
        payload = request.get_json(force=True)
        return jsonify(pm.save_pipeline(payload))

    @app.route("/api/pipelines/<string:name>/yaml", methods=["PUT"])
    def save_pipeline_yaml(name: str) -> Response:
        """Save YAML text directly to file.

        Args:
            name: Pipeline name

        Returns:
            JSON response with save result
        """
        yaml_content = request.get_data(as_text=True)
        return jsonify(pm.save_pipeline_yaml(name, yaml_content))

    @app.route("/api/pipelines/parse", methods=["POST"])
    def parse_pipeline_yaml() -> Response:
        """Parse arbitrary YAML text and return structured data.

        Used for frontend validation/synchronization with canvas.

        Returns:
            JSON response with parsed pipeline structure and raw YAML
        """
        yaml_content = request.get_data(as_text=True)
        parsed = pm.parse_pipeline_yaml_content(yaml_content)

        # Also return original text for frontend consistency
        if isinstance(parsed, dict):
            parsed["_raw_yaml"] = yaml_content
        return jsonify(parsed)

    @app.route("/api/pipelines/<string:name>", methods=["GET"])
    def get_pipeline(name: str):
        return jsonify(pm.load_pipeline(name))

    @app.route("/api/pipelines/<string:name>", methods=["DELETE"])
    def delete_pipeline(name: str):
        pm.delete_pipeline(name)
        return jsonify({"status": "deleted"})

    @app.route("/api/pipelines/<string:name>/rename", methods=["POST"])
    def rename_pipeline(name: str):
        """Rename a pipeline"""
        payload = request.get_json(force=True)
        new_name = payload.get("new_name", "").strip()
        if not new_name:
            return jsonify({"error": "new_name is required"}), 400

        result = pm.rename_pipeline(name, new_name)
        return jsonify(result)

    @app.route("/api/pipelines/<string:name>/parameters", methods=["GET"])
    def get_parameters(name: str):
        return jsonify(pm.load_parameters(name))

    @app.route("/api/pipelines/<string:name>/parameters", methods=["PUT"])
    def save_parameters(name: str):
        payload = request.get_json(force=True)
        pm.save_parameters(name, payload)
        return jsonify({"status": "saved"})

    @app.route("/api/pipelines/<string:name>/build", methods=["POST"])
    def build_pipeline(name: str):
        return jsonify(pm.build(name))

    @app.route("/api/pipelines/<string:name>/demo/start", methods=["POST"])
    def start_demo_session(name: str) -> Response:
        """Start a demo session for a pipeline.

        Args:
            name: Pipeline name

        Returns:
            JSON response with session info
        """
        payload = request.get_json(force=True) or {}
        session_id = payload.get("session_id")
        if not session_id:
            return jsonify({"error": "session_id is required"}), 400
        return jsonify(pm.start_demo_session(name, session_id))

    @app.route("/api/pipelines/demo/stop", methods=["POST"])
    def stop_demo_session() -> Response:
        """Stop a demo session.

        Returns:
            JSON response with stop result
        """
        payload = request.get_json(force=True) or {}
        session_id = payload.get("session_id")
        if not session_id:
            return jsonify({"error": "session_id is required"}), 400
        return jsonify(pm.stop_demo_session(session_id))

    @app.route("/api/pipelines/<string:name>/chat", methods=["POST"])
    def chat_pipeline(name: str) -> Response:
        """Handle chat request for a pipeline.

        Args:
            name: Pipeline name

        Returns:
            Server-sent events stream response
        """
        payload = request.get_json(force=True)
        question = payload.get("question", "")
        session_id = payload.get("session_id")
        dynamic_params = payload.get("dynamic_params", {})

        # New: Frontend-provided conversation history (previous conversations in browser session)
        # Compatible with two field names: conversation_history or history
        raw_history = (
            payload.get("conversation_history") or payload.get("history") or []
        )

        # Convert format: frontend uses {role, text}, backend uses {role, content}
        conversation_history = []
        for msg in raw_history:
            if isinstance(msg, dict):
                role = msg.get("role", "user")
                # Compatible with both text and content field names
                content = msg.get("content") or msg.get("text") or ""
                if role in ("user", "assistant") and content:
                    conversation_history.append({"role": role, "content": str(content)})

        # New: Allow frontend to force full pipeline execution
        force_full_pipeline = payload.get("force_full_pipeline", False)

        selected_collection = dynamic_params.get("collection_name")

        try:
            kb_config = pm.load_kb_config()
            milvus_global_config = kb_config.get("milvus", {})

            retriever_params = {
                "index_backend": "milvus",
                "index_backend_configs": {"milvus": milvus_global_config},
            }

            if selected_collection:
                retriever_params["collection_name"] = selected_collection
                LOGGER.debug(f"Chat using collection override: {selected_collection}")

            dynamic_params["retriever"] = retriever_params

            if "collection_name" in dynamic_params:
                del dynamic_params["collection_name"]

        except Exception as e:
            LOGGER.warning(f"Failed to construct retriever config: {e}")

        if not session_id:
            return (
                jsonify({"error": "session_id missing. Please start engine first."}),
                400,
            )

        # Determine if this is the first question based on frontend-provided conversation history
        # Note: Frontend has already added the current question to history before sending request
        # So: First question has 1 message in history, second question has 3 messages in history
        # Only when there are >= 2 messages (indicating previous assistant reply) do we use multi-turn
        has_previous_conversation = len(conversation_history) >= 2

        LOGGER.info(
            f"Session {session_id}: conversation_history length = "
            f"{len(conversation_history)}, has_previous = {has_previous_conversation}"
        )

        if not has_previous_conversation or force_full_pipeline:
            # First question (or forced): run full pipeline
            LOGGER.info(
                f"Session {session_id}: First turn (history empty), "
                f"running full pipeline '{name}'"
            )
            return Response(
                pm.chat_demo_stream(name, question, session_id, dynamic_params),
                mimetype="text/event-stream",
            )
        else:
            # Subsequent questions: use multi-turn chat with frontend-provided history
            LOGGER.info(
                f"Session {session_id}: Multi-turn chat mode with "
                f"{len(conversation_history)} history messages"
            )
            return Response(
                pm.chat_multiturn_stream(
                    session_id, question, dynamic_params, conversation_history
                ),
                mimetype="text/event-stream",
            )

    @app.route("/api/pipelines/chat/stop", methods=["POST"])
    def stop_chat_generation() -> Response:
        """Stop chat generation for a session.

        Returns:
            JSON response with stop result
        """
        payload = request.get_json(force=True) or {}
        session_id = payload.get("session_id")
        if not session_id:
            return jsonify({"error": "session_id required"}), 400
        return jsonify(pm.interrupt_chat(session_id))

    @app.route("/api/pipelines/chat/clear-history", methods=["POST"])
    def clear_chat_history() -> Response:
        """Clear conversation history for a session.

        This makes the next question run the full pipeline again.

        Returns:
            JSON response with clear status
        """
        payload = request.get_json(force=True) or {}
        session_id = payload.get("session_id")
        if not session_id:
            return jsonify({"error": "session_id required"}), 400

        session = pm.SESSION_MANAGER.get(session_id)
        if not session:
            return jsonify({"error": "Session not found"}), 404

        session.clear_history()
        return jsonify(
            {
                "status": "cleared",
                "session_id": session_id,
                "message": (
                    "Conversation history cleared. "
                    "Next question will run full pipeline."
                ),
            }
        )

    @app.route("/api/pipelines/chat/history", methods=["GET"])
    def get_chat_history() -> Response:
        """Get conversation history for current session.

        Returns:
            JSON response with conversation history
        """
        session_id = request.args.get("session_id")
        if not session_id:
            return jsonify({"error": "session_id required"}), 400

        # Validate session_id format to prevent injection attacks
        import re
        if not re.match(r'^[a-zA-Z0-9_-]+$', session_id) or len(session_id) > 128:
            return jsonify({"error": "Invalid session_id format"}), 400

        session = pm.SESSION_MANAGER.get(session_id)
        if not session:
            return jsonify({"error": "Session not found"}), 404

        # Verify session ownership by checking client IP or authentication token
        # This prevents session hijacking by validating the request context
        client_ip = request.remote_addr
        if hasattr(session, 'client_ip') and session.client_ip != client_ip:
            return jsonify({"error": "Session validation failed"}), 403

        history = session.get_conversation_history()
        return jsonify(
            {
                "session_id": session_id,
                "history": history,
                "is_first_turn": session.is_first_turn(),
                "message_count": len(history),
            }
        )

    @app.route("/api/chat/export/docx", methods=["POST"])
    def export_chat_docx() -> Response:
        """Export chat content to a DOCX file."""
        payload = request.get_json(force=True) or {}
        answer_text = payload.get("text", "")
        if not isinstance(answer_text, str) or not answer_text.strip():
            return jsonify({"error": "text is required"}), 400

        question_text = payload.get("question", "")
        sources = payload.get("sources", [])
        if not isinstance(sources, list):
            sources = []

        try:
            docx_bytes, filename = _build_chat_export_docx(
                question_text, answer_text, sources
            )
        except Exception as e:
            LOGGER.error("DOCX export failed: %s", e, exc_info=True)
            return jsonify({"error": str(e)}), 500

        return Response(
            docx_bytes,
            mimetype=DOCX_MIME_TYPE,
            headers={"Content-Disposition": _build_content_disposition(filename)},
        )

    # ===== Background Chat Task API =====

    @app.route("/api/pipelines/<string:name>/chat/background", methods=["POST"])
    def start_background_chat(name: str) -> Response:
        """Start a background chat task.

        Args:
            name: Pipeline name

        Returns:
            JSON response with task ID
        """
        payload = request.get_json(force=True)
        question = payload.get("question", "")
        session_id = payload.get("session_id")
        dynamic_params = payload.get("dynamic_params", {})
        user_id = payload.get("user_id", "") or request.headers.get("X-User-ID", "")

        if not question:
            return jsonify({"error": "question is required"}), 400
        if not session_id:
            return (
                jsonify(
                    {"error": "session_id is required. Please start engine first."}
                ),
                400,
            )

        # Handle collection_name
        selected_collection = dynamic_params.get("collection_name")
        try:
            kb_config = pm.load_kb_config()
            milvus_global_config = kb_config.get("milvus", {})

            retriever_params = {
                "index_backend": "milvus",
                "index_backend_configs": {"milvus": milvus_global_config},
            }

            if selected_collection:
                retriever_params["collection_name"] = selected_collection

            dynamic_params["retriever"] = retriever_params

            if "collection_name" in dynamic_params:
                del dynamic_params["collection_name"]

        except Exception as e:
            LOGGER.warning(f"Failed to construct retriever config: {e}")

        try:
            task_id = pm.run_background_chat(
                name, question, session_id, dynamic_params, user_id=user_id
            )
            return (
                jsonify(
                    {
                        "status": "started",
                        "task_id": task_id,
                        "message": "Task started in background",
                    }
                ),
                202,
            )
        except Exception as e:
            LOGGER.error(f"Failed to start background chat: {e}")
            return jsonify({"error": str(e)}), 500

    @app.route("/api/background-tasks", methods=["GET"])
    def list_background_tasks() -> Response:
        """List background tasks for current user.

        Returns:
            JSON response with list of tasks
        """
        limit = request.args.get("limit", 20, type=int)
        user_id = request.args.get("user_id", "") or request.headers.get(
            "X-User-ID", ""
        )
        LOGGER.info(f"Listing background tasks for user_id: '{user_id}'")
        tasks = pm.list_background_tasks(limit, user_id=user_id)
        LOGGER.info(f"Found {len(tasks)} tasks for user_id: '{user_id}'")
        return jsonify(tasks)

    @app.route("/api/background-tasks/<string:task_id>", methods=["GET"])
    def get_background_task(task_id: str) -> Response:
        """Get single background task status.

        Args:
            task_id: Task identifier

        Returns:
            JSON response with task information
        """
        user_id = request.args.get("user_id", "") or request.headers.get(
            "X-User-ID", ""
        )
        task = pm.get_background_task(task_id, user_id=user_id)
        if not task:
            return jsonify({"error": "Task not found"}), 404
        return jsonify(task)

    @app.route("/api/background-tasks/<string:task_id>", methods=["DELETE"])
    def delete_background_task(task_id: str) -> Response:
        """Delete a background task.

        Args:
            task_id: Task identifier

        Returns:
            JSON response with deletion status
        """
        user_id = request.args.get("user_id", "") or request.headers.get(
            "X-User-ID", ""
        )
        success = pm.delete_background_task(task_id, user_id=user_id)
        if success:
            return jsonify({"status": "deleted", "task_id": task_id})
        return jsonify({"error": "Task not found"}), 404

    @app.route("/api/background-tasks/clear-completed", methods=["POST"])
    def clear_completed_tasks() -> Response:
        """Clear completed background tasks for current user.

        Returns:
            JSON response with clear status and count
        """
        payload = request.get_json(force=True) if request.is_json else {}
        user_id = payload.get("user_id", "") or request.headers.get("X-User-ID", "")
        count = pm.clear_completed_background_tasks(user_id=user_id)
        return jsonify({"status": "cleared", "count": count})

    @app.route("/api/system/shutdown", methods=["POST"])
    def shutdown() -> Response:
        """Shutdown the server.

        Returns:
            JSON response with shutdown status
        """
        LOGGER.info("Shutdown requested")
        func = request.environ.get("werkzeug.server.shutdown")
        if func:
            threading.Timer(0.2, func).start()
            return jsonify({"status": "shutting-down", "mode": "graceful"})

        threading.Timer(0.5, os._exit, args=(0,)).start()
        return jsonify({"status": "shutting-down", "mode": "force"})

    @app.route("/api/kb/config", methods=["GET"])
    def get_kb_config() -> Response:
        """Get knowledge base configuration.

        Returns:
            JSON response with KB configuration
        """
        return jsonify(pm.load_kb_config())

    @app.route("/api/kb/config", methods=["POST"])
    def save_kb_config() -> Response:
        """Save knowledge base configuration.

        Returns:
            JSON response with save status
        """
        try:
            payload = request.get_json(force=True)
            pm.save_kb_config(payload)
            return jsonify({"status": "saved"})
        except Exception as e:
            return jsonify({"error": str(e)}), 500

    @app.route("/api/kb/files", methods=["GET"])
    def list_kb_files() -> Response:
        """List knowledge base files.

        Returns:
            JSON response with list of KB files
        """
        return jsonify(pm.list_kb_files())

    @app.route("/api/kb/files/inspect", methods=["GET"])
    def inspect_kb_folder() -> Response:
        """Inspect files in a knowledge base folder.

        Returns:
            JSON response with list of files in folder
        """
        category = request.args.get("category", "raw")
        folder_name = request.args.get("name")

        if not folder_name:
            return jsonify({"error": "Folder name required"}), 400

        base_dir = {
            "raw": pm.KB_RAW_DIR,
            "corpus": pm.KB_CORPUS_DIR,
            "chunks": pm.KB_CHUNKS_DIR,
        }.get(category)

        target_path = base_dir / folder_name

        if not target_path.exists() or not target_path.is_dir():
            return jsonify({"error": "Folder not found"}), 404

        files = []
        for f in sorted(target_path.glob("*")):
            if f.is_file() and not f.name.startswith("."):
                files.append({"name": f.name, "size": f.stat().st_size})

        return jsonify({"files": files})

    @app.route("/api/kb/upload", methods=["POST"])
    def upload_kb_file() -> Response:
        """Upload files to knowledge base.

        Returns:
            JSON response with upload result
        """
        files = request.files.getlist("file")

        if not files or files[0].filename == "":
            return jsonify({"error": "No selected files"}), 400

        try:
            result = pm.upload_kb_files_batch(files)
            return jsonify(result)
        except Exception as e:
            LOGGER.error(f"Upload failed: {e}")
            return jsonify({"error": str(e)}), 500

    @app.route("/api/kb/files/<string:category>/<string:filename>", methods=["DELETE"])
    def delete_kb_file(category: str, filename: str) -> Response:
        """Delete a knowledge base file.

        Args:
            category: File category (raw, corpus, chunks, collection, index)
            filename: File name

        Returns:
            JSON response with deletion result
        """
        if category not in ["raw", "corpus", "chunks", "collection", "index"]:
            return jsonify({"error": "Invalid category"}), 400

        try:
            return jsonify(pm.delete_kb_file(category, filename))
        except Exception as e:
            return jsonify({"error": str(e)}), 500

    @app.route("/api/kb/staging/clear", methods=["POST"])
    def clear_staging_area() -> Response:
        """Clear staging area: delete all files in raw, corpus, chunks directories.

        Returns:
            JSON response with clear result
        """
        try:
            result = pm.clear_staging_area()
            return jsonify(result)
        except Exception as e:
            LOGGER.error(f"Failed to clear staging area: {e}")
            return jsonify({"error": str(e)}), 500

    @app.route("/api/kb/run", methods=["POST"])
    def run_kb_task() -> Response:
        """Run a knowledge base pipeline task.

        Returns:
            JSON response with task submission status
        """
        payload = request.get_json(force=True)
        pipeline_name = payload.get("pipeline_name")
        target_file = payload.get("target_file")

        collection_name = payload.get("collection_name")
        index_mode = payload.get("index_mode", "append")

        chunk_params = {
            "chunk_backend": payload.get("chunk_backend", "token"),
            "tokenizer_or_token_counter": payload.get(
                "tokenizer_or_token_counter", "gpt2"
            ),
            "chunk_size": payload.get("chunk_size", 500),
            "use_title": payload.get("use_title", True),
        }

        # Embedding parameters (for milvus_index)
        embedding_params = {
            "api_key": payload.get("emb_api_key", ""),
            "base_url": payload.get("emb_base_url", "https://api.openai.com/v1"),
            "model_name": payload.get("emb_model_name", "text-embedding-3-small"),
        }

        if not pipeline_name or not target_file:
            return jsonify({"error": "Missing pipeline_name or target_file"}), 400

        output_dir = ""
        if pipeline_name == "build_text_corpus":
            output_dir = str(pm.KB_CORPUS_DIR)
        elif pipeline_name == "corpus_chunk":
            output_dir = str(pm.KB_CHUNKS_DIR)
        elif pipeline_name == "milvus_index":
            output_dir = ""

        task_id = str(uuid.uuid4())
        KB_TASKS[task_id] = {
            "status": "running",
            "pipeline": pipeline_name,
            "created_at": datetime.now().isoformat(),
        }

        thread = threading.Thread(
            target=_run_kb_background,
            args=(
                task_id,
                pipeline_name,
                target_file,
                output_dir,
                collection_name,
                index_mode,
                chunk_params,
                embedding_params,
            ),
            daemon=True,  # Set as daemon thread to auto-exit when main program exits, preventing hangs
        )
        thread.start()

        return (
            jsonify(
                {
                    "status": "submitted",
                    "task_id": task_id,
                    "message": "Task started in background",
                }
            ),
            202,
        )

    @app.route("/api/kb/status/<string:task_id>", methods=["GET"])
    def get_kb_task_status(task_id: str) -> Response:
        """Get knowledge base task status.

        Args:
            task_id: Task identifier

        Returns:
            JSON response with task status
        """
        task = KB_TASKS.get(task_id)
        if not task:
            return jsonify({"error": "Task not found"}), 404

        return jsonify(task)

    # =========================================
    # Prompt Template API
    # =========================================

    PROMPTS_DIR = BASE_DIR.parent.parent / "prompt"

    @app.route("/api/prompts", methods=["GET"])
    def list_prompts() -> Response:
        """List all prompt template files.

        Returns:
            JSON response with list of prompt files
        """
        prompts = []
        if PROMPTS_DIR.exists():
            for f in sorted(PROMPTS_DIR.rglob("*.jinja*")):
                rel_path = f.relative_to(PROMPTS_DIR)
                prompts.append(
                    {"name": f.name, "path": str(rel_path), "size": f.stat().st_size}
                )
        return jsonify(prompts)

    @app.route("/api/prompts/<path:filepath>", methods=["GET"])
    def get_prompt(filepath: str) -> Response:
        """Read a prompt template file.

        Args:
            filepath: Path to prompt file

        Returns:
            JSON response with file content
        """
        file_path = PROMPTS_DIR / filepath
        if not file_path.exists():
            return jsonify({"error": "File not found"}), 404
        if not str(file_path.resolve()).startswith(str(PROMPTS_DIR.resolve())):
            return jsonify({"error": "Invalid path"}), 400

        content = file_path.read_text(encoding="utf-8")
        return jsonify({"path": filepath, "content": content})

    @app.route("/api/prompts", methods=["POST"])
    def create_prompt() -> Response:
        """Create a new prompt template file.

        Returns:
            JSON response with creation status
        """
        payload = request.get_json(force=True)
        name = payload.get("name", "").strip()
        content = payload.get("content", "")

        if not name:
            return jsonify({"error": "Name is required"}), 400

        # Security check: ensure filename is valid
        if ".." in name or name.startswith("/"):
            return jsonify({"error": "Invalid filename"}), 400

        file_path = PROMPTS_DIR / name
        if file_path.exists():
            return jsonify({"error": "File already exists"}), 409

        # Ensure directory exists
        file_path.parent.mkdir(parents=True, exist_ok=True)
        file_path.write_text(content, encoding="utf-8")

        return jsonify({"status": "created", "path": name})

    @app.route("/api/prompts/<path:filepath>", methods=["PUT"])
    def update_prompt(filepath: str):
        """Update a prompt template file"""
        payload = request.get_json(force=True)
        content = payload.get("content", "")

        file_path = PROMPTS_DIR / filepath
        if not str(file_path.resolve()).startswith(str(PROMPTS_DIR.resolve())):
            return jsonify({"error": "Invalid path"}), 400

        # Ensure directory exists
        file_path.parent.mkdir(parents=True, exist_ok=True)
        file_path.write_text(content, encoding="utf-8")

        return jsonify({"status": "saved", "path": filepath})

    @app.route("/api/prompts/<path:filepath>", methods=["DELETE"])
    def delete_prompt(filepath: str):
        """Delete a prompt template file"""
        file_path = PROMPTS_DIR / filepath
        if not file_path.exists():
            return jsonify({"error": "File not found"}), 404
        if not str(file_path.resolve()).startswith(str(PROMPTS_DIR.resolve())):
            return jsonify({"error": "Invalid path"}), 400

        file_path.unlink()
        return jsonify({"status": "deleted"})

    @app.route("/api/prompts/<path:filepath>/rename", methods=["POST"])
    def rename_prompt(filepath: str):
        """Rename a prompt template file"""
        payload = request.get_json(force=True)
        new_name = payload.get("new_name", "").strip()

        if not new_name:
            return jsonify({"error": "new_name is required"}), 400

        if ".." in new_name or new_name.startswith("/"):
            return jsonify({"error": "Invalid filename"}), 400

        old_path = PROMPTS_DIR / filepath
        if not old_path.exists():
            return jsonify({"error": "File not found"}), 404
        if not str(old_path.resolve()).startswith(str(PROMPTS_DIR.resolve())):
            return jsonify({"error": "Invalid path"}), 400

        new_path = PROMPTS_DIR / new_name
        if new_path.exists():
            return jsonify({"error": "A file with this name already exists"}), 409

        # Ensure target directory exists
        new_path.parent.mkdir(parents=True, exist_ok=True)
        old_path.rename(new_path)

        return jsonify(
            {"status": "renamed", "old_path": filepath, "new_path": new_name}
        )

    # =========================================
    # AI Assistant API
    # =========================================

    @app.route("/api/ai/test", methods=["POST"])
    def test_ai_connection():
        """Test AI API connection"""
        import requests

        payload = request.get_json(force=True)
        provider = payload.get("provider", "openai")
        base_url = payload.get("baseUrl", "").rstrip("/")
        api_key = payload.get("apiKey", "")
        model = payload.get("model", "gpt-5-mini")

        if not api_key:
            return jsonify({"success": False, "error": "API key is required"})

        try:
            if provider == "openai" or provider == "custom":
                # OpenAI-compatible API
                headers = {
                    "Authorization": f"Bearer {api_key}",
                    "Content-Type": "application/json",
                }

                # Try to list models or send a simple request
                test_url = f"{base_url}/models"
                resp = requests.get(test_url, headers=headers, timeout=10)

                if resp.status_code == 200:
                    return jsonify({"success": True, "model": model})
                else:
                    # Try a simple completion as fallback
                    chat_url = f"{base_url}/chat/completions"
                    test_payload = {
                        "model": model,
                        "messages": [{"role": "user", "content": "Hi"}],
                        "max_tokens": 5,
                    }
                    resp = requests.post(
                        chat_url, headers=headers, json=test_payload, timeout=15
                    )

                    if resp.status_code == 200:
                        return jsonify({"success": True, "model": model})
                    else:
                        return jsonify(
                            {
                                "success": False,
                                "error": f"API returned {resp.status_code}: {resp.text[:200]}",
                            }
                        )

            elif provider == "anthropic":
                headers = {
                    "x-api-key": api_key,
                    "anthropic-version": "2023-06-01",
                    "Content-Type": "application/json",
                }
                test_url = f"{base_url}/messages"
                test_payload = {
                    "model": model,
                    "max_tokens": 5,
                    "messages": [{"role": "user", "content": "Hi"}],
                }
                resp = requests.post(
                    test_url, headers=headers, json=test_payload, timeout=15
                )

                if resp.status_code == 200:
                    return jsonify({"success": True, "model": model})
                else:
                    return jsonify(
                        {
                            "success": False,
                            "error": f"API returned {resp.status_code}: {resp.text[:200]}",
                        }
                    )

            elif provider == "azure":
                headers = {"api-key": api_key, "Content-Type": "application/json"}
                # Azure uses deployment name in URL
                test_url = f"{base_url}/openai/deployments/{model}/chat/completions?api-version=2024-02-15-preview"
                test_payload = {
                    "messages": [{"role": "user", "content": "Hi"}],
                    "max_tokens": 5,
                }
                resp = requests.post(
                    test_url, headers=headers, json=test_payload, timeout=15
                )

                if resp.status_code == 200:
                    return jsonify({"success": True, "model": model})
                else:
                    return jsonify(
                        {
                            "success": False,
                            "error": f"API returned {resp.status_code}: {resp.text[:200]}",
                        }
                    )
            else:
                return jsonify(
                    {"success": False, "error": f"Unknown provider: {provider}"}
                )

        except requests.Timeout:
            return jsonify({"success": False, "error": "Connection timeout"})
        except requests.RequestException as e:
            return jsonify({"success": False, "error": str(e)})
        except Exception as e:
            return jsonify({"success": False, "error": str(e)})

    @app.route("/api/ai/chat", methods=["POST"])
    def ai_chat():
        """Handle AI chat request"""
        import requests

        def safe_first_item(value: Any) -> Dict[str, Any]:
            if isinstance(value, list) and value:
                first = value[0]
                return first if isinstance(first, dict) else {}
            return {}

        def safe_json_response(resp: requests.Response) -> Dict[str, Any]:
            try:
                return json.loads(resp.content)
            except Exception:
                try:
                    return resp.json()
                except Exception:
                    return {}

        def decode_sse_line(raw_line: Any) -> str:
            if isinstance(raw_line, bytes):
                return raw_line.decode("utf-8", errors="replace")
            return str(raw_line)

        payload = request.get_json(force=True)
        settings = payload.get("settings", {})
        messages = payload.get("messages", [])
        context = payload.get("context", {})
        stream_response = payload.get("stream", False)

        provider = settings.get("provider", "openai")
        base_url = settings.get("baseUrl", "").rstrip("/")
        api_key = settings.get("apiKey", "")
        model = settings.get("model", "gpt-5-mini")

        if not api_key:
            return jsonify({"error": "API key is required"})

        # Build system prompt with context
        system_prompt = build_ai_system_prompt(context)

        # Prepend system message
        full_messages = [{"role": "system", "content": system_prompt}] + messages

        try:
            if stream_response and (provider == "openai" or provider == "custom"):

                def stream_openai_chat():
                    import requests

                    try:
                        headers = {
                            "Authorization": f"Bearer {api_key}",
                            "Content-Type": "application/json",
                        }
                        chat_url = f"{base_url}/chat/completions"
                        chat_payload = {
                            "model": model,
                            "messages": full_messages,
                            "temperature": 0.7,
                            "stream": True,
                        }
                        resp = requests.post(
                            chat_url,
                            headers=headers,
                            json=chat_payload,
                            stream=True,
                            timeout=120,
                        )
                        resp.encoding = "utf-8"

                        if resp.status_code != 200:
                            yield f"data: {json.dumps({'type': 'error', 'message': f'API error: {resp.status_code}'})}\n\n"
                            return

                        full_text = ""
                        for raw_line in resp.iter_lines(decode_unicode=False):
                            if not raw_line:
                                continue

                            line = decode_sse_line(raw_line).strip()
                            if not line:
                                continue

                            if line.startswith("data:"):
                                line = line[len("data:") :].strip()

                            if line == "[DONE]":
                                break

                            try:
                                data = json.loads(line)
                            except Exception:
                                continue

                            choice = safe_first_item(data.get("choices"))
                            delta_block = choice.get("delta")
                            if not isinstance(delta_block, dict):
                                delta_block = {}
                            delta = delta_block.get("content", "")
                            if delta:
                                full_text += delta
                                yield f"data: {json.dumps({'type': 'token', 'content': delta, 'is_final': False})}\n\n"

                        actions = parse_ai_actions(full_text, context)
                        yield f"data: {json.dumps({'type': 'final', 'content': full_text, 'actions': actions})}\n\n"
                    except requests.Timeout:
                        yield f"data: {json.dumps({'type': 'error', 'message': 'Request timeout'})}\n\n"
                    except Exception as e:
                        LOGGER.error(f"AI chat stream error: {e}")
                        yield f"data: {json.dumps({'type': 'error', 'message': str(e)})}\n\n"

                stream_headers = {
                    "Cache-Control": "no-cache",
                    "Connection": "keep-alive",
                    "X-Accel-Buffering": "no",
                }
                return Response(
                    stream_with_context(stream_openai_chat()),
                    content_type="text/event-stream; charset=utf-8",
                    headers=stream_headers,
                )

            if provider == "openai" or provider == "custom":
                headers = {
                    "Authorization": f"Bearer {api_key}",
                    "Content-Type": "application/json",
                }
                chat_url = f"{base_url}/chat/completions"
                chat_payload = {
                    "model": model,
                    "messages": full_messages,
                    "temperature": 0.7,
                }
                resp = requests.post(
                    chat_url, headers=headers, json=chat_payload, timeout=120
                )
                resp.encoding = "utf-8"

                if resp.status_code == 200:
                    data = safe_json_response(resp)
                    choice = safe_first_item(data.get("choices"))
                    message_block = choice.get("message")
                    if not isinstance(message_block, dict):
                        message_block = {}
                    content = message_block.get("content", "")

                    # Parse for actions
                    actions = parse_ai_actions(content, context)

                    return jsonify({"content": content, "actions": actions})
                else:
                    return jsonify({"error": f"API error: {resp.status_code}"})

            elif provider == "anthropic":
                headers = {
                    "x-api-key": api_key,
                    "anthropic-version": "2023-06-01",
                    "Content-Type": "application/json",
                }

                # Convert messages format for Anthropic
                anthropic_messages = []
                for msg in messages:
                    anthropic_messages.append(
                        {"role": msg["role"], "content": msg["content"]}
                    )

                chat_url = f"{base_url}/messages"
                chat_payload = {
                    "model": model,
                    "max_tokens": 4096,
                    "system": system_prompt,
                    "messages": anthropic_messages,
                }
                resp = requests.post(
                    chat_url, headers=headers, json=chat_payload, timeout=120
                )
                resp.encoding = "utf-8"

                if resp.status_code == 200:
                    data = safe_json_response(resp)
                    content_entries = data.get("content")
                    if isinstance(content_entries, list):
                        content = safe_first_item(content_entries).get("text", "")
                    elif isinstance(content_entries, str):
                        content = content_entries
                    else:
                        content = ""
                    actions = parse_ai_actions(content, context)
                    return jsonify({"content": content, "actions": actions})
                else:
                    return jsonify({"error": f"API error: {resp.status_code}"})

            elif provider == "azure":
                headers = {"api-key": api_key, "Content-Type": "application/json"}
                chat_url = f"{base_url}/openai/deployments/{model}/chat/completions?api-version=2024-02-15-preview"
                chat_payload = {"messages": full_messages, "temperature": 0.7}
                resp = requests.post(
                    chat_url, headers=headers, json=chat_payload, timeout=120
                )
                resp.encoding = "utf-8"

                if resp.status_code == 200:
                    data = safe_json_response(resp)
                    choice = safe_first_item(data.get("choices"))
                    message_block = choice.get("message")
                    if not isinstance(message_block, dict):
                        message_block = {}
                    content = message_block.get("content", "")
                    actions = parse_ai_actions(content, context)
                    return jsonify({"content": content, "actions": actions})
                else:
                    return jsonify({"error": f"API error: {resp.status_code}"})
            else:
                return jsonify({"error": f"Unknown provider: {provider}"})

        except requests.Timeout:
            return jsonify({"error": "Request timeout"})
        except Exception as e:
            LOGGER.error(f"AI chat error: {e}")
            return jsonify({"error": str(e)})

    return app


def build_ai_system_prompt(context: Dict) -> str:
    """Build system prompt with current context"""
    base_prompt = """You are an AI assistant for UltraRAG, a RAG (Retrieval-Augmented Generation) pipeline configuration system.

You help users:
1. Build and configure pipelines (YAML format)
2. Set parameters for servers (retriever, generator, reranker, etc.)
3. Edit Jinja2 prompt templates

When suggesting modifications, use the following format to indicate actionable changes:

For Pipeline modifications:
```yaml:pipeline
<your yaml content here>
```

For Prompt modifications:
```jinja:prompt:<filename>
<your jinja content here>
```

For Parameter changes, describe them clearly with the full path like:
"Set `generation.model_name` to `gpt-5-mini`"

Be concise and helpful. Provide complete code when suggesting changes.
"""

    # Add context information
    context_info = []

    if context.get("currentMode"):
        context_info.append(f"Current mode: {context['currentMode']}")

    if context.get("selectedPipeline"):
        context_info.append(f"Selected pipeline: {context['selectedPipeline']}")

    if context.get("pipelineYaml"):
        context_info.append(
            f"Current pipeline YAML:\n```yaml\n{context['pipelineYaml']}\n```"
        )

    if context.get("currentPromptFile"):
        context_info.append(f"Current prompt file: {context['currentPromptFile']}")
        if context.get("promptContent"):
            context_info.append(
                f"Current prompt content:\n```jinja\n{context['promptContent']}\n```"
            )

    if context.get("parameters"):
        import json

        params_str = json.dumps(context["parameters"], indent=2, ensure_ascii=False)
        context_info.append(f"Current parameters:\n```json\n{params_str}\n```")

    if context_info:
        base_prompt += "\n\n## Current Context\n" + "\n".join(context_info)

    llms_doc = load_llms_doc()
    if llms_doc:
        base_prompt += "\n\n## Repository Reference (llms.txt)\n" + llms_doc

    return base_prompt


def deduplicate_ai_actions(actions: list[dict]) -> list[dict]:
    """Remove duplicated apply actions while preserving order."""
    seen = set()
    unique = []

    for action in actions:
        try:
            key = json.dumps(
                {
                    "type": action.get("type"),
                    "filename": action.get("filename"),
                    "path": action.get("path"),
                    "preview": action.get("preview"),
                    "content": action.get("content"),
                    "value": action.get("value"),
                },
                sort_keys=True,
                ensure_ascii=False,
            )
        except Exception:
            key = f"{action.get('type')}::{action.get('filename') or action.get('path')}::{action.get('preview') or action.get('content')}"

        if key in seen:
            continue
        seen.add(key)
        unique.append(action)

    return unique


def parse_ai_actions(content: str, context: Dict) -> list:
    """Parse AI response for actionable modifications"""
    import re

    actions = []

    # Parse pipeline YAML blocks
    yaml_pattern = r"```yaml:pipeline\s*\n(.*?)```"
    yaml_matches = re.findall(yaml_pattern, content, re.DOTALL)
    for match in yaml_matches:
        actions.append(
            {
                "type": "modify_pipeline",
                "content": match.strip(),
                "preview": match.strip()[:500],
            }
        )

    # Parse prompt blocks
    prompt_pattern = r"```jinja:prompt:([^\n]+)\s*\n(.*?)```"
    prompt_matches = re.findall(prompt_pattern, content, re.DOTALL)
    for filename, prompt_content in prompt_matches:
        actions.append(
            {
                "type": "modify_prompt",
                "filename": filename.strip(),
                "content": prompt_content.strip(),
                "preview": prompt_content.strip()[:500],
            }
        )

    # Parse parameter changes
    param_pattern = r"[Ss]et\s+`([^`]+)`\s+to\s+`([^`]+)`"
    param_matches = re.findall(param_pattern, content)
    for path, value in param_matches:
        # Try to parse value
        try:
            import json

            parsed_value = json.loads(value)
        except Exception:
            parsed_value = value

        actions.append(
            {
                "type": "modify_parameter",
                "path": path.strip(),
                "value": parsed_value,
                "preview": f"{path} = {value}",
            }
        )

    return deduplicate_ai_actions(actions)


if __name__ == "__main__":
    import os
    
    logging.basicConfig(level=logging.INFO)
    app = create_app()
    
    # Security: Use environment variables to control debug mode and host
    # Never enable debug mode in production or expose debugger to network
    debug_mode = os.getenv("FLASK_DEBUG", "False").lower() == "true"
    host = os.getenv("FLASK_HOST", "127.0.0.1" if debug_mode else "0.0.0.0")
    port = int(os.getenv("FLASK_PORT", "5050"))
    
    # Additional safety: Never allow debug=True with host='0.0.0.0'
    if debug_mode and host == "0.0.0.0":
        app.logger.warning(
            "Security warning: Debug mode should not be enabled with host='0.0.0.0'. "
            "Using host='127.0.0.1' instead."
        )
        host = "127.0.0.1"
    
    app.run(host=host, port=port, debug=debug_mode)
